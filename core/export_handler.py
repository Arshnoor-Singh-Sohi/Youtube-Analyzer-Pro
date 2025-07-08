# core/export_handler.py
import json
import io
from datetime import datetime
from typing import Dict, Any
import streamlit as st

# For PDF generation
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# For Word document generation
try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

class ExportHandler:
    """Handle various export formats for analysis results"""
    
    def __init__(self):
        self.pdf_available = PDF_AVAILABLE
        self.docx_available = DOCX_AVAILABLE
        
        if not self.pdf_available:
            st.warning("PDF export not available. Install reportlab: pip install reportlab")
        if not self.docx_available:
            st.warning("Word export not available. Install python-docx: pip install python-docx")
    
    def export_summary(self, analysis_results: Dict[str, Any], format_type: str) -> bytes:
        """Export just the summary in specified format"""
        
        if format_type.upper() == "PDF" and self.pdf_available:
            return self._create_summary_pdf(analysis_results)
        elif format_type.upper() == "WORD DOCUMENT" and self.docx_available:
            return self._create_summary_docx(analysis_results)
        elif format_type.upper() == "JSON":
            return self._create_summary_json(analysis_results)
        else:
            return self._create_summary_txt(analysis_results)
    
    def export_full_report(self, analysis_results: Dict[str, Any], transcript_data: Dict, video_info: Dict, format_type: str) -> bytes:
        """Export comprehensive report in specified format"""
        
        if format_type.upper() == "PDF" and self.pdf_available:
            return self._create_full_report_pdf(analysis_results, transcript_data, video_info)
        elif format_type.upper() == "WORD DOCUMENT" and self.docx_available:
            return self._create_full_report_docx(analysis_results, transcript_data, video_info)
        elif format_type.upper() == "JSON":
            return self._create_full_report_json(analysis_results, transcript_data, video_info)
        else:
            return self._create_full_report_txt(analysis_results, transcript_data, video_info)
    
    def get_mime_type(self, format_type: str) -> str:
        """Get MIME type for download"""
        format_type = format_type.upper()
        
        mime_types = {
            "PDF": "application/pdf",
            "WORD DOCUMENT": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "JSON": "application/json",
            "TEXT FILE": "text/plain"
        }
        
        return mime_types.get(format_type, "text/plain")
    
    def _create_summary_txt(self, analysis_results: Dict[str, Any]) -> bytes:
        """Create text summary"""
        content = []
        content.append("=" * 60)
        content.append("VIDEO ANALYSIS SUMMARY")
        content.append("=" * 60)
        content.append(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        content.append("")
        
        # Main Summary
        content.append("MAIN SUMMARY")
        content.append("-" * 20)
        content.append(analysis_results.get('main_summary', 'Not available'))
        content.append("")
        
        # Key Takeaways
        if 'key_takeaways' in analysis_results:
            content.append("KEY TAKEAWAYS")
            content.append("-" * 20)
            for i, takeaway in enumerate(analysis_results['key_takeaways'], 1):
                content.append(f"{i}. {takeaway}")
            content.append("")
        
        # Important Quotes
        if 'important_quotes' in analysis_results:
            content.append("IMPORTANT QUOTES")
            content.append("-" * 20)
            for quote in analysis_results['important_quotes']:
                content.append(f'• "{quote}"')
            content.append("")
        
        # Action Items
        if 'action_items' in analysis_results:
            content.append("ACTION ITEMS")
            content.append("-" * 20)
            for item in analysis_results['action_items']:
                content.append(f"• {item}")
            content.append("")
        
        return "\n".join(content).encode('utf-8')
    
    def _create_summary_json(self, analysis_results: Dict[str, Any]) -> bytes:
        """Create JSON summary"""
        summary_data = {
            "export_info": {
                "type": "summary",
                "generated_at": datetime.now().isoformat(),
                "version": "1.0"
            },
            "main_summary": analysis_results.get('main_summary', ''),
            "key_takeaways": analysis_results.get('key_takeaways', []),
            "important_quotes": analysis_results.get('important_quotes', []),
            "action_items": analysis_results.get('action_items', []),
            "topics": analysis_results.get('topics', []),
            "sentiment_analysis": analysis_results.get('sentiment_analysis', {})
        }
        
        return json.dumps(summary_data, indent=2, ensure_ascii=False).encode('utf-8')
    
    def _create_summary_pdf(self, analysis_results: Dict[str, Any]) -> bytes:
        """Create PDF summary"""
        if not self.pdf_available:
            return self._create_summary_txt(analysis_results)
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1  # Center alignment
        )
        story.append(Paragraph("Video Analysis Summary", title_style))
        story.append(Spacer(1, 12))
        
        # Generated info
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        story.append(Spacer(1, 20))
        
        # Main Summary
        story.append(Paragraph("Main Summary", styles['Heading2']))
        story.append(Paragraph(analysis_results.get('main_summary', 'Not available'), styles['Normal']))
        story.append(Spacer(1, 15))
        
        # Key Takeaways
        if 'key_takeaways' in analysis_results:
            story.append(Paragraph("Key Takeaways", styles['Heading2']))
            for i, takeaway in enumerate(analysis_results['key_takeaways'], 1):
                story.append(Paragraph(f"{i}. {takeaway}", styles['Normal']))
            story.append(Spacer(1, 15))
        
        # Important Quotes
        if 'important_quotes' in analysis_results:
            story.append(Paragraph("Important Quotes", styles['Heading2']))
            for quote in analysis_results['important_quotes']:
                story.append(Paragraph(f'"{quote}"', styles['BodyText']))
            story.append(Spacer(1, 15))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.read()
    
    def _create_summary_docx(self, analysis_results: Dict[str, Any]) -> bytes:
        """Create Word document summary"""
        if not self.docx_available:
            return self._create_summary_txt(analysis_results)
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Video Analysis Summary', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Generated info
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Main Summary
        doc.add_heading('Main Summary', level=1)
        doc.add_paragraph(analysis_results.get('main_summary', 'Not available'))
        
        # Key Takeaways
        if 'key_takeaways' in analysis_results:
            doc.add_heading('Key Takeaways', level=1)
            for i, takeaway in enumerate(analysis_results['key_takeaways'], 1):
                doc.add_paragraph(f"{i}. {takeaway}", style='List Number')
        
        # Important Quotes
        if 'important_quotes' in analysis_results:
            doc.add_heading('Important Quotes', level=1)
            for quote in analysis_results['important_quotes']:
                p = doc.add_paragraph(f'"{quote}"')
                p.style = 'Intense Quote'
        
        # Action Items
        if 'action_items' in analysis_results:
            doc.add_heading('Action Items', level=1)
            for item in analysis_results['action_items']:
                doc.add_paragraph(item, style='List Bullet')
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    def _create_full_report_txt(self, analysis_results: Dict[str, Any], transcript_data: Dict, video_info: Dict) -> bytes:
        """Create comprehensive text report"""
        content = []
        content.append("=" * 80)
        content.append("COMPREHENSIVE VIDEO ANALYSIS REPORT")
        content.append("=" * 80)
        content.append(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        content.append("")
        
        # Video Information
        content.append("VIDEO INFORMATION")
        content.append("-" * 30)
        content.append(f"Title: {video_info.get('title', 'Unknown')}")
        content.append(f"Channel: {video_info.get('channel', 'Unknown')}")
        content.append(f"Duration: {video_info.get('duration', 'Unknown')}")
        content.append(f"URL: {video_info.get('url', 'Unknown')}")
        content.append("")
        
        # Main Summary
        content.append("EXECUTIVE SUMMARY")
        content.append("-" * 30)
        content.append(analysis_results.get('main_summary', 'Not available'))
        content.append("")
        
        # Key Takeaways
        if 'key_takeaways' in analysis_results:
            content.append("KEY TAKEAWAYS")
            content.append("-" * 30)
            for i, takeaway in enumerate(analysis_results['key_takeaways'], 1):
                content.append(f"{i}. {takeaway}")
            content.append("")
        
        # Topics
        if 'topics' in analysis_results:
            content.append("MAIN TOPICS DISCUSSED")
            content.append("-" * 30)
            content.append(", ".join(analysis_results['topics']))
            content.append("")
        
        # Sentiment Analysis
        if 'sentiment_analysis' in analysis_results:
            sentiment = analysis_results['sentiment_analysis']
            content.append("SENTIMENT ANALYSIS")
            content.append("-" * 30)
            content.append(f"Positive: {sentiment.get('positive', 0):.1%}")
            content.append(f"Neutral: {sentiment.get('neutral', 0):.1%}")
            content.append(f"Negative: {sentiment.get('negative', 0):.1%}")
            content.append(f"Overall Score: {sentiment.get('overall_score', 0):.2f}")
            content.append("")
        
        # Important Quotes
        if 'important_quotes' in analysis_results:
            content.append("IMPORTANT QUOTES")
            content.append("-" * 30)
            for quote in analysis_results['important_quotes']:
                content.append(f'• "{quote}"')
            content.append("")
        
        # Action Items
        if 'action_items' in analysis_results:
            content.append("ACTION ITEMS")
            content.append("-" * 30)
            for item in analysis_results['action_items']:
                content.append(f"• {item}")
            content.append("")
        
        # Timeline
        if 'timeline' in analysis_results:
            content.append("VIDEO TIMELINE")
            content.append("-" * 30)
            for event in analysis_results['timeline']:
                content.append(f"{event['timestamp']}: {event['description']}")
            content.append("")
        
        # Q&A
        if 'questions_and_answers' in analysis_results:
            content.append("QUESTIONS & ANSWERS")
            content.append("-" * 30)
            for qa in analysis_results['questions_and_answers']:
                content.append(f"Q: {qa['question']}")
                content.append(f"A: {qa['answer']}")
                content.append("")
        
        # Statistics
        content.append("STATISTICS")
        content.append("-" * 30)
        content.append(f"Total Words: {len(transcript_data.get('text', '').split()):,}")
        content.append(f"Total Segments: {transcript_data.get('total_segments', 0):,}")
        content.append(f"Estimated Reading Time: {len(transcript_data.get('text', '').split()) // 200 + 1} minutes")
        content.append("")
        
        return "\n".join(content).encode('utf-8')
    
    def _create_full_report_json(self, analysis_results: Dict[str, Any], transcript_data: Dict, video_info: Dict) -> bytes:
        """Create comprehensive JSON report"""
        report_data = {
            "export_info": {
                "type": "full_report",
                "generated_at": datetime.now().isoformat(),
                "version": "1.0"
            },
            "video_info": video_info,
            "analysis_results": analysis_results,
            "transcript_stats": {
                "total_words": len(transcript_data.get('text', '').split()),
                "total_segments": transcript_data.get('total_segments', 0),
                "estimated_reading_time_minutes": len(transcript_data.get('text', '').split()) // 200 + 1,
                "language_codes": transcript_data.get('language_codes', [])
            }
        }
        
        return json.dumps(report_data, indent=2, ensure_ascii=False).encode('utf-8')
    
    def _create_full_report_pdf(self, analysis_results: Dict[str, Any], transcript_data: Dict, video_info: Dict) -> bytes:
        """Create comprehensive PDF report"""
        if not self.pdf_available:
            return self._create_full_report_txt(analysis_results, transcript_data, video_info)
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # Title page
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=1
        )
        story.append(Paragraph("Comprehensive Video Analysis Report", title_style))
        story.append(Spacer(1, 50))
        
        # Video info table
        video_data = [
            ['Video Information', ''],
            ['Title', video_info.get('title', 'Unknown')],
            ['Channel', video_info.get('channel', 'Unknown')],
            ['Duration', video_info.get('duration', 'Unknown')],
            ['Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S')]
        ]
        
        video_table = Table(video_data, colWidths=[2*inch, 4*inch])
        video_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(video_table)
        story.append(Spacer(1, 30))
        
        # Sections
        sections = [
            ('Executive Summary', analysis_results.get('main_summary', '')),
            ('Key Takeaways', '\n'.join([f"• {item}" for item in analysis_results.get('key_takeaways', [])])),
            ('Important Quotes', '\n'.join([f'"{quote}"' for quote in analysis_results.get('important_quotes', [])])),
            ('Action Items', '\n'.join([f"• {item}" for item in analysis_results.get('action_items', [])]))
        ]
        
        for section_title, section_content in sections:
            if section_content:
                story.append(Paragraph(section_title, styles['Heading2']))
                story.append(Paragraph(section_content, styles['Normal']))
                story.append(Spacer(1, 15))
        
        doc.build(story)
        buffer.seek(0)
        return buffer.read()
    
    def _create_full_report_docx(self, analysis_results: Dict[str, Any], transcript_data: Dict, video_info: Dict) -> bytes:
        """Create comprehensive Word document report"""
        if not self.docx_available:
            return self._create_full_report_txt(analysis_results, transcript_data, video_info)
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Comprehensive Video Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Video Information
        doc.add_heading('Video Information', level=1)
        info_table = doc.add_table(rows=5, cols=2)
        info_table.style = 'Table Grid'
        
        info_data = [
            ('Title', video_info.get('title', 'Unknown')),
            ('Channel', video_info.get('channel', 'Unknown')),
            ('Duration', video_info.get('duration', 'Unknown')),
            ('URL', video_info.get('url', 'Unknown')),
            ('Generated', datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = value
        
        # Analysis sections
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(analysis_results.get('main_summary', 'Not available'))
        
        if 'key_takeaways' in analysis_results:
            doc.add_heading('Key Takeaways', level=1)
            for takeaway in analysis_results['key_takeaways']:
                doc.add_paragraph(takeaway, style='List Bullet')
        
        if 'important_quotes' in analysis_results:
            doc.add_heading('Important Quotes', level=1)
            for quote in analysis_results['important_quotes']:
                p = doc.add_paragraph(f'"{quote}"')
                p.style = 'Intense Quote'
        
        if 'action_items' in analysis_results:
            doc.add_heading('Action Items', level=1)
            for item in analysis_results['action_items']:
                doc.add_paragraph(item, style='List Bullet')
        
        # Statistics
        doc.add_heading('Statistics', level=1)
        stats_table = doc.add_table(rows=3, cols=2)
        stats_table.style = 'Table Grid'
        
        stats_data = [
            ('Total Words', f"{len(transcript_data.get('text', '').split()):,}"),
            ('Total Segments', f"{transcript_data.get('total_segments', 0):,}"),
            ('Reading Time', f"{len(transcript_data.get('text', '').split()) // 200 + 1} minutes")
        ]
        
        for i, (label, value) in enumerate(stats_data):
            stats_table.cell(i, 0).text = label
            stats_table.cell(i, 1).text = value
        
        # Save to buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()